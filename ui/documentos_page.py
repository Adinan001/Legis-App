# ui/documentos_page.py — Etapa 3: editor completo, importação Word, variáveis automáticas
import os
import subprocess
from datetime import date
from PyQt6.QtWidgets import (QWidget, QVBoxLayout, QHBoxLayout, QLabel,
                             QTableWidget, QTableWidgetItem, QHeaderView,
                             QAbstractItemView, QDialog, QFormLayout,
                             QComboBox, QTextEdit, QMessageBox, QDialogButtonBox,
                             QSplitter, QFileDialog, QToolButton, QMenu,
                             QFrame, QSizePolicy, QInputDialog, QLineEdit,
                             QPushButton)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QFont, QAction, QTextCharFormat, QColor, QKeySequence, QShortcut
from config import COLORS
from ui.widgets import (btn_primario, btn_secundario, btn_perigo, campo_busca,
                        input_field, label_titulo, estilo_tabela, separador_h,
                        estilo_dialog)
from core.database import (buscar_documentos, salvar_documento, atualizar_documento,
                           excluir_documento, toggle_favorito, buscar_variaveis_processo,
                           migrar_etapa3)

CATEGORIAS = ["Cível", "Penal", "Trabalhista", "Tributário", "Contratos",
              "Petições", "Recursos", "Pareceres", "Administrativo",
              "Procurações", "Habeas Corpus", "Mandado de Segurança", "Outro"]

# Variáveis disponíveis para inserção nos documentos
VARIAVEIS = {
    "{{numero_processo}}":   "Número do processo",
    "{{cliente}}":           "Nome do cliente",
    "{{acao}}":              "Área / tipo de ação",
    "{{status}}":            "Status do processo",
    "{{data_distribuicao}}": "Data de distribuição",
    "{{cpf_cnpj}}":          "CPF ou CNPJ do cliente",
    "{{contato_cliente}}":   "Telefone do cliente",
    "{{email_cliente}}":     "E-mail do cliente",
    "{{data_hoje}}":         "Data de hoje",
    "{{ano_atual}}":         "Ano atual",
}


# ──────────────────────────────────────────────
# Thread de exportação
# ──────────────────────────────────────────────
class ExportadorDoc(QThread):
    concluido = pyqtSignal(str)
    erro      = pyqtSignal(str)

    def __init__(self, formato, caminho, titulo, categoria, conteudo):
        super().__init__()
        self.formato = formato; self.caminho = caminho
        self.titulo = titulo; self.categoria = categoria; self.conteudo = conteudo

    def run(self):
        try:
            import tempfile, shutil, os
            # Gera primeiro num arquivo temporário (fora do OneDrive),
            # depois move para o destino — evita crash de sincronização
            ext = ".docx" if self.formato == "docx" else ".pdf"
            fd, tmp_path = tempfile.mkstemp(suffix=ext)
            os.close(fd)
            self._caminho_real = self.caminho
            self.caminho = tmp_path

            if self.formato == "docx":
                self._exportar_docx()
            else:
                self._exportar_pdf()

            # Mover para o destino final
            destino = self._caminho_real
            os.makedirs(os.path.dirname(destino), exist_ok=True)
            shutil.move(tmp_path, destino)
            self.caminho = destino
            self.concluido.emit(destino)
        except Exception as e:
            import traceback
            self.erro.emit(traceback.format_exc())

    def _exportar_docx(self):
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(3); section.bottom_margin = Cm(2)
            section.left_margin = Cm(3); section.right_margin = Cm(2)
        cab = doc.add_paragraph()
        cab.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = cab.add_run(f"{self.categoria.upper()}  |  {date.today().strftime('%d/%m/%Y')}")
        r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x5A, 0x6B, 0x5C)
        doc.add_paragraph("─" * 80)
        tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = tp.add_run(self.titulo.upper())
        rt.font.size = Pt(14); rt.font.bold = True
        doc.add_paragraph()
        import re
        for linha in self.conteudo.split("\n"):
            # Limpar markdown residual (###, ---, ===)
            linha = re.sub(r"^#{1,6}\s*", "", linha)  # remove ### no início
            if linha.strip() in ("---", "***", "___", "===") or re.fullmatch(r"[-*_=]{3,}", linha.strip() or "x"):
                continue  # pula linhas separadoras
            par = doc.add_paragraph()
            linha_strip = linha.strip()
            # Linha inteiramente em negrito/título (começa e tradicionalmente é cabeçalho de seção)
            eh_titulo = linha_strip.startswith("**") and linha_strip.endswith("**") and linha_strip.count("**") == 2
            if eh_titulo:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER if any(
                    k in linha_strip.upper() for k in ["EXCELENT", "EGRÉGIO", "COLENDO", "SENHOR"]) else WD_ALIGN_PARAGRAPH.LEFT
            else:
                par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            par.paragraph_format.first_line_indent = Cm(2) if not eh_titulo and linha_strip else None
            # Interpretar **negrito** dentro da linha
            partes = re.split(r"(\*\*.*?\*\*)", linha)
            for parte in partes:
                if parte.startswith("**") and parte.endswith("**") and len(parte) > 4:
                    run = par.add_run(parte[2:-2])
                    run.font.bold = True
                else:
                    run = par.add_run(parte)
                run.font.size = Pt(12); run.font.name = "Times New Roman"
        doc.add_paragraph()
        rod = doc.add_paragraph(); rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = rod.add_run("Gerado pelo Legis Beta — Sistema de Gestão Jurídica")
        rr.font.size = Pt(8); rr.font.color.rgb = RGBColor(0x8A, 0x9B, 0x8C)
        doc.save(self.caminho)

    def _exportar_pdf(self):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors as rc
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        doc = SimpleDocTemplate(self.caminho, pagesize=A4,
                                rightMargin=2*cm, leftMargin=3*cm,
                                topMargin=3*cm, bottomMargin=2*cm)
        verde = rc.HexColor("#3A6B40"); cinza = rc.HexColor("#5A6B5C"); escuro = rc.HexColor("#1A2B1C")
        el = []
        el.append(Paragraph(f"{self.categoria.upper()}  |  {date.today().strftime('%d/%m/%Y')}",
            ParagraphStyle("c", fontName="Helvetica", fontSize=8, textColor=cinza, alignment=2, spaceAfter=4)))
        el.append(HRFlowable(width="100%", thickness=1.5, color=verde, spaceAfter=12))
        el.append(Paragraph(self.titulo.upper(),
            ParagraphStyle("t", fontName="Helvetica-Bold", fontSize=14,
                           textColor=escuro, alignment=1, spaceAfter=16)))
        el.append(Spacer(1, 0.3*cm))
        body = ParagraphStyle("b", fontName="Times-Roman", fontSize=12,
                               textColor=escuro, alignment=4, leading=18, spaceAfter=4)
        for linha in self.conteudo.split("\n"):
            el.append(Paragraph(linha.strip() or " ", body))
        el.append(Spacer(1, 1*cm))
        el.append(HRFlowable(width="100%", thickness=0.5, color=cinza))
        el.append(Paragraph("Gerado pelo Legis Beta — Sistema de Gestão Jurídica",
            ParagraphStyle("r", fontName="Helvetica", fontSize=8, textColor=cinza, alignment=1)))
        doc.build(el)


# ──────────────────────────────────────────────
# Thread de importação Word
# ──────────────────────────────────────────────
class ImportadorWord(QThread):
    concluido = pyqtSignal(str, str)  # titulo, conteudo
    erro      = pyqtSignal(str)

    def __init__(self, caminho):
        super().__init__()
        self.caminho = caminho

    def run(self):
        try:
            from docx import Document
            doc = Document(self.caminho)
            linhas = []
            for par in doc.paragraphs:
                linhas.append(par.text)
            conteudo = "\n".join(linhas)
            titulo = os.path.splitext(os.path.basename(self.caminho))[0]
            self.concluido.emit(titulo, conteudo)
        except Exception as e:
            import traceback
            self.erro.emit(traceback.format_exc())


# ──────────────────────────────────────────────
# Dialog novo documento
# ──────────────────────────────────────────────
class NovoDocumentoDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Novo Documento")
        self.setMinimumWidth(420)
        self.setStyleSheet(estilo_dialog())
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(14)
        titulo_lbl = QLabel("Novo Documento")
        titulo_lbl.setFont(QFont("Segoe UI", 14, QFont.Weight.Bold))
        titulo_lbl.setStyleSheet(f"color: {COLORS['text_primary']};")
        layout.addWidget(titulo_lbl)
        layout.addWidget(separador_h())
        form = QFormLayout(); form.setSpacing(12)
        self.txt_titulo = input_field("Ex: Procuração Ad Judicia, Petição Inicial")
        self.cb_cat = QComboBox()
        self.cb_cat.addItems(CATEGORIAS)
        self.cb_cat.setStyleSheet(f"padding: 8px; border: 1.5px solid {COLORS['border']}; border-radius: 6px; background: white; font-size: 12px;")
        def lbl(t):
            l = QLabel(f"<b>{t}</b>"); l.setStyleSheet(f"color: {COLORS['text_primary']}; font-size: 12px; border: none;"); return l
        form.addRow(lbl("Título:"), self.txt_titulo)
        form.addRow(lbl("Categoria:"), self.cb_cat)
        layout.addLayout(form)
        botoes = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        botoes.button(QDialogButtonBox.StandardButton.Ok).setText("Criar")
        botoes.button(QDialogButtonBox.StandardButton.Cancel).setText("Cancelar")
        botoes.button(QDialogButtonBox.StandardButton.Ok).setStyleSheet(f"background-color: {COLORS['accent']}; color: white; font-weight: 600; padding: 8px 20px; border-radius: 6px; border: none;")
        botoes.button(QDialogButtonBox.StandardButton.Cancel).setStyleSheet(f"background-color: white; color: {COLORS['text_secondary']}; font-weight: 600; padding: 8px 20px; border-radius: 6px; border: 1.5px solid {COLORS['border']};")
        botoes.accepted.connect(self.validar); botoes.rejected.connect(self.reject)
        layout.addWidget(botoes)

    def validar(self):
        if not self.txt_titulo.text().strip():
            QMessageBox.warning(self, "Campo obrigatório", "Informe o título."); return
        self.accept()

    def obter_dados(self):
        return {"titulo": self.txt_titulo.text().strip(),
                "categoria": self.cb_cat.currentText(), "conteudo": ""}


# ──────────────────────────────────────────────
# Barra de formatação do editor
# ──────────────────────────────────────────────
class BarraFormatacao(QWidget):
    def __init__(self, editor, parent=None):
        super().__init__(parent)
        self.editor = editor
        self.setStyleSheet(f"background: {COLORS['white']}; border-bottom: 1px solid {COLORS['border']};")
        lay = QHBoxLayout(self)
        lay.setContentsMargins(8, 4, 8, 4)
        lay.setSpacing(4)

        def btn_fmt(texto, tooltip, callback):
            b = QPushButton(texto)
            b.setToolTip(tooltip)
            b.setFixedSize(30, 28)
            b.setCursor(Qt.CursorShape.PointingHandCursor)
            b.setStyleSheet(f"""
                QPushButton {{
                    background: transparent; border: 1px solid transparent;
                    border-radius: 4px; font-size: 13px; font-weight: bold;
                    color: {COLORS['text_primary']};
                }}
                QPushButton:hover {{ background: {COLORS['accent_light']}; border-color: {COLORS['accent']}; }}
                QPushButton:pressed {{ background: {COLORS['accent']}22; }}
            """)
            b.clicked.connect(callback)
            return b

        lay.addWidget(btn_fmt("B", "Negrito (Ctrl+B)", self._negrito))
        lay.addWidget(btn_fmt("I", "Itálico (Ctrl+I)", self._italico))
        lay.addWidget(btn_fmt("U", "Sublinhado (Ctrl+U)", self._sublinhado))

        # Separador
        sep = QFrame(); sep.setFrameShape(QFrame.Shape.VLine)
        sep.setStyleSheet(f"color: {COLORS['border']};"); sep.setFixedWidth(1)
        lay.addWidget(sep)

        # Tamanho da fonte
        self.cb_tamanho = QComboBox()
        self.cb_tamanho.addItems(["10", "11", "12", "13", "14", "16", "18", "20", "24"])
        self.cb_tamanho.setCurrentText("12")
        self.cb_tamanho.setFixedWidth(60)
        self.cb_tamanho.setStyleSheet(f"padding: 3px; border: 1px solid {COLORS['border']}; border-radius: 4px; font-size: 11px;")
        self.cb_tamanho.currentTextChanged.connect(self._tamanho_fonte)
        lay.addWidget(QLabel("pt"))
        lay.addWidget(self.cb_tamanho)

        sep2 = QFrame(); sep2.setFrameShape(QFrame.Shape.VLine)
        sep2.setStyleSheet(f"color: {COLORS['border']};"); sep2.setFixedWidth(1)
        lay.addWidget(sep2)

        # Alinhamentos
        lay.addWidget(btn_fmt("≡", "Justificado", self._alinhar_justificado))
        lay.addWidget(btn_fmt("⬛", "Esquerda", self._alinhar_esquerda))
        lay.addWidget(btn_fmt("▪", "Centro", self._alinhar_centro))
        lay.addWidget(btn_fmt("▫", "Direita", self._alinhar_direita))

        sep3 = QFrame(); sep3.setFrameShape(QFrame.Shape.VLine)
        sep3.setStyleSheet(f"color: {COLORS['border']};"); sep3.setFixedWidth(1)
        lay.addWidget(sep3)

        # Inserir variável
        btn_var = QPushButton("{ } Variável")
        btn_var.setToolTip("Inserir variável automática")
        btn_var.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_var.setStyleSheet(f"""
            QPushButton {{
                background: {COLORS['accent_light']}; color: {COLORS['accent']};
                border: 1px solid {COLORS['accent']}50; border-radius: 5px;
                font-size: 11px; font-weight: 600; padding: 4px 10px;
            }}
            QPushButton:hover {{ background: {COLORS['accent']}20; }}
        """)
        btn_var.clicked.connect(self._menu_variaveis)
        lay.addWidget(btn_var)
        lay.addStretch()

    def _negrito(self):
        fmt = QTextCharFormat()
        cursor = self.editor.textCursor()
        fmt.setFontWeight(QFont.Weight.Bold if cursor.charFormat().fontWeight() != 700 else QFont.Weight.Normal)
        cursor.mergeCharFormat(fmt)

    def _italico(self):
        fmt = QTextCharFormat()
        cursor = self.editor.textCursor()
        fmt.setFontItalic(not cursor.charFormat().fontItalic())
        cursor.mergeCharFormat(fmt)

    def _sublinhado(self):
        fmt = QTextCharFormat()
        cursor = self.editor.textCursor()
        fmt.setFontUnderline(not cursor.charFormat().fontUnderline())
        cursor.mergeCharFormat(fmt)

    def _tamanho_fonte(self, tamanho):
        try:
            fmt = QTextCharFormat()
            fmt.setFontPointSize(float(tamanho))
            self.editor.textCursor().mergeCharFormat(fmt)
            self.editor.setFocus()
        except Exception:
            pass

    def _alinhar_justificado(self):
        self.editor.setAlignment(Qt.AlignmentFlag.AlignJustify)
    def _alinhar_esquerda(self):
        self.editor.setAlignment(Qt.AlignmentFlag.AlignLeft)
    def _alinhar_centro(self):
        self.editor.setAlignment(Qt.AlignmentFlag.AlignHCenter)
    def _alinhar_direita(self):
        self.editor.setAlignment(Qt.AlignmentFlag.AlignRight)

    def _menu_variaveis(self):
        menu = QMenu(self)
        menu.setStyleSheet(f"""
            QMenu {{
                background: {COLORS['white']}; border: 1px solid {COLORS['border']};
                border-radius: 6px; padding: 4px; font-size: 12px;
            }}
            QMenu::item {{ padding: 7px 16px; border-radius: 4px; color: {COLORS['text_primary']}; }}
            QMenu::item:selected {{ background: {COLORS['accent_light']}; color: {COLORS['accent']}; }}
            QMenu::separator {{ height: 1px; background: {COLORS['border']}; margin: 4px 8px; }}
        """)
        for var, descricao in VARIAVEIS.items():
            act = QAction(f"{var}   —   {descricao}", self)
            act.triggered.connect(lambda _, v=var: self.editor.insertPlainText(v))
            menu.addAction(act)
        menu.exec(self.mapToGlobal(self.rect().bottomLeft()))


# ──────────────────────────────────────────────
# Dialog para preencher variáveis
# ──────────────────────────────────────────────
class PreencherVariaveisDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Preencher Variáveis Automaticamente")
        self.setMinimumWidth(440)
        self.setStyleSheet(estilo_dialog())
        layout = QVBoxLayout(self)
        layout.setContentsMargins(24, 20, 24, 20)
        layout.setSpacing(14)

        titulo = QLabel("Preencher Variáveis do Documento")
        titulo.setFont(QFont("Segoe UI", 13, QFont.Weight.Bold))
        titulo.setStyleSheet(f"color: {COLORS['text_primary']};")
        layout.addWidget(titulo)

        desc = QLabel("Informe o número do processo para substituir automaticamente\nas variáveis {{nome_cliente}}, {{numero_processo}}, etc.")
        desc.setStyleSheet(f"color: {COLORS['text_secondary']}; font-size: 11px;")
        desc.setWordWrap(True)
        layout.addWidget(desc)
        layout.addWidget(separador_h())

        form = QFormLayout(); form.setSpacing(10)
        self.txt_processo = input_field("Ex: 1002345-67.2026.8.26.0000")
        def lbl(t):
            l = QLabel(f"<b>{t}</b>"); l.setStyleSheet(f"color: {COLORS['text_primary']}; font-size: 12px; border: none;"); return l
        form.addRow(lbl("Nº do Processo:"), self.txt_processo)
        layout.addLayout(form)

        # Preview das variáveis que serão substituídas
        self.lbl_preview = QLabel("")
        self.lbl_preview.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 10px; border: none;")
        self.lbl_preview.setWordWrap(True)
        layout.addWidget(self.lbl_preview)
        self.txt_processo.textChanged.connect(self._preview_variaveis)

        botoes = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        botoes.button(QDialogButtonBox.StandardButton.Ok).setText("Preencher")
        botoes.button(QDialogButtonBox.StandardButton.Cancel).setText("Cancelar")
        botoes.button(QDialogButtonBox.StandardButton.Ok).setStyleSheet(f"background-color: {COLORS['accent']}; color: white; font-weight: 600; padding: 8px 20px; border-radius: 6px; border: none;")
        botoes.button(QDialogButtonBox.StandardButton.Cancel).setStyleSheet(f"background-color: white; color: {COLORS['text_secondary']}; font-weight: 600; padding: 8px 20px; border-radius: 6px; border: 1.5px solid {COLORS['border']};")
        botoes.accepted.connect(self.accept)
        botoes.rejected.connect(self.reject)
        layout.addWidget(botoes)

    def _preview_variaveis(self, numero):
        if len(numero) < 5: self.lbl_preview.setText(""); return
        vars_ = buscar_variaveis_processo(numero)
        if vars_:
            txt = "  •  ".join([f"{k.replace('{{','').replace('}}','')}: {v}" for k, v in list(vars_.items())[:4] if v])
            self.lbl_preview.setText(f"✅  {txt}")
        else:
            self.lbl_preview.setText("⚠️  Processo não encontrado no banco.")

    def obter_numero(self):
        return self.txt_processo.text().strip()


# ──────────────────────────────────────────────
# Página principal
# ──────────────────────────────────────────────


class ExportadorThread(QThread):
    """Thread leve que delega a exportação ao subprocess (isola python-docx do Qt)."""
    concluido = pyqtSignal(str)
    erro = pyqtSignal(str)

    def __init__(self, formato, caminho, titulo, categoria, conteudo):
        super().__init__()
        self.formato = formato
        self.caminho = caminho
        self.titulo = titulo
        self.categoria = categoria
        self.conteudo = conteudo

    def run(self):
        try:
            from core.ai.runner import exportar_documento
            ok, resultado = exportar_documento(
                self.formato, self.caminho, self.titulo, self.categoria, self.conteudo)
            if ok:
                self.concluido.emit(resultado)
            else:
                self.erro.emit(resultado)
        except Exception as e:
            import traceback
            self.erro.emit(traceback.format_exc())


class DocumentosPage(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self._dados = []
        self._doc_aberto = None
        self._exportador = None
        self._importador = None
        self._modificado = False
        try: migrar_etapa3()
        except Exception: pass
        self.init_ui()

    def init_ui(self):
        raiz = QHBoxLayout(self)
        raiz.setContentsMargins(0, 0, 0, 0)
        raiz.setSpacing(0)

        # ── PAINEL ESQUERDO — lista ──
        painel_lista = QWidget()
        painel_lista.setFixedWidth(290)
        painel_lista.setStyleSheet(f"background: {COLORS['white']}; border-right: 1px solid {COLORS['border']};")
        lay_lista = QVBoxLayout(painel_lista)
        lay_lista.setContentsMargins(14, 16, 14, 14)
        lay_lista.setSpacing(8)

        lbl_docs = QLabel("Documentos")
        lbl_docs.setFont(QFont("Segoe UI", 13, QFont.Weight.Bold))
        lbl_docs.setStyleSheet(f"color: {COLORS['text_primary']};")
        lay_lista.addWidget(lbl_docs)

        self.busca = campo_busca("Buscar...")
        self.busca.textChanged.connect(self.filtrar)
        lay_lista.addWidget(self.busca)

        # Botões de ação na lista
        btn_novo = btn_primario("＋  Novo")
        btn_novo.clicked.connect(self.abrir_novo)
        lay_lista.addWidget(btn_novo)

        # Gerar com IA
        btn_gerar_ia = QPushButton("🤖  Gerar com IA")
        btn_gerar_ia.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_gerar_ia.setStyleSheet(f"""
            QPushButton {{
                background: {COLORS['accent']}; color: white;
                border: none; border-radius: 8px; padding: 10px;
                font-size: 12px; font-weight: 700;
            }}
            QPushButton:hover {{ background: {COLORS['accent_hover']}; }}
        """)
        btn_gerar_ia.clicked.connect(self.gerar_com_ia)
        lay_lista.addWidget(btn_gerar_ia)

        # Importar Word
        btn_importar = btn_secundario("📥  Importar Word (.docx)")
        btn_importar.clicked.connect(self.importar_word)
        lay_lista.addWidget(btn_importar)

        lay_lista.addWidget(separador_h())

        # Filtro por categoria
        self.cb_filtro_cat = QComboBox()
        self.cb_filtro_cat.addItem("Todas as categorias", "")
        for cat in CATEGORIAS:
            self.cb_filtro_cat.addItem(cat, cat)
        self.cb_filtro_cat.setStyleSheet(f"padding: 6px; border: 1px solid {COLORS['border']}; border-radius: 6px; font-size: 11px; background: white;")
        self.cb_filtro_cat.currentIndexChanged.connect(self.filtrar)
        lay_lista.addWidget(self.cb_filtro_cat)

        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(["Título", "Cat."])
        self.table.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        self.table.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SelectionMode.SingleSelection)
        self.table.setStyleSheet(f"""
            QTableWidget {{ background: transparent; border: none; font-size: 12px; color: {COLORS['text_primary']}; }}
            QTableWidget::item {{ padding: 8px 6px; border-radius: 4px; }}
            QTableWidget::item:selected {{ background: {COLORS['accent_light']}; color: {COLORS['accent']}; }}
            QHeaderView::section {{ background: {COLORS['white']}; padding: 6px; border: none;
                border-bottom: 1px solid {COLORS['border']}; font-weight: 700; font-size: 10px;
                color: {COLORS['text_muted']}; text-transform: uppercase; }}
        """)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.table.verticalHeader().setVisible(False)
        self.table.setShowGrid(False)
        self.table.clicked.connect(self.abrir_documento)
        lay_lista.addWidget(self.table)

        self.btn_excluir = btn_perigo("🗑  Excluir")
        self.btn_excluir.clicked.connect(self.excluir_selecionado)
        lay_lista.addWidget(self.btn_excluir)

        raiz.addWidget(painel_lista)

        # ── PAINEL DIREITO — editor ──
        painel_editor = QWidget()
        painel_editor.setStyleSheet(f"background: {COLORS['bg_main']};")
        lay_editor = QVBoxLayout(painel_editor)
        lay_editor.setContentsMargins(0, 0, 0, 0)
        lay_editor.setSpacing(0)

        # Barra superior do editor
        barra_topo = QWidget()
        barra_topo.setStyleSheet(f"background: {COLORS['white']}; border-bottom: 1px solid {COLORS['border']};")
        barra_topo_lay = QHBoxLayout(barra_topo)
        barra_topo_lay.setContentsMargins(16, 10, 16, 10)
        barra_topo_lay.setSpacing(10)

        self.lbl_doc_aberto = QLabel("Selecione ou crie um documento")
        self.lbl_doc_aberto.setFont(QFont("Segoe UI", 13, QFont.Weight.Bold))
        self.lbl_doc_aberto.setStyleSheet(f"color: {COLORS['text_primary']}; border: none;")
        barra_topo_lay.addWidget(self.lbl_doc_aberto)

        self.lbl_status_salvo = QLabel("")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 11px; border: none;")
        barra_topo_lay.addWidget(self.lbl_status_salvo)
        barra_topo_lay.addStretch()

        # Botão preencher variáveis
        self.btn_variaveis = btn_secundario("⚡  Preencher Variáveis")
        self.btn_variaveis.setEnabled(False)
        self.btn_variaveis.clicked.connect(self.preencher_variaveis)
        barra_topo_lay.addWidget(self.btn_variaveis)

        # Botão avaliar (aprendizado da IA)
        self.btn_avaliar = btn_secundario("⭐  Avaliar")
        self.btn_avaliar.setEnabled(False)
        self.btn_avaliar.clicked.connect(self.avaliar_peca)
        barra_topo_lay.addWidget(self.btn_avaliar)

        # Salvar
        self.btn_salvar = btn_primario("💾  Salvar")
        self.btn_salvar.setEnabled(False)
        self.btn_salvar.clicked.connect(self.salvar_edicao)
        barra_topo_lay.addWidget(self.btn_salvar)

        # Exportar
        self.btn_exportar = QToolButton()
        self.btn_exportar.setText("📤  Exportar ▾")
        self.btn_exportar.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self.btn_exportar.setEnabled(False)
        self.btn_exportar.setStyleSheet(f"""
            QToolButton {{ background: {COLORS['white']}; color: {COLORS['text_primary']};
                font-weight: 600; font-size: 12px; padding: 8px 14px; border-radius: 6px;
                border: 1.5px solid {COLORS['border']}; }}
            QToolButton:hover {{ background: {COLORS['accent_light']}; border-color: {COLORS['accent']}; color: {COLORS['accent']}; }}
            QToolButton:disabled {{ color: {COLORS['text_muted']}; }}
            QToolButton::menu-indicator {{ image: none; }}
        """)
        menu_exp = QMenu(self)
        menu_exp.setStyleSheet(f"QMenu {{ background: {COLORS['white']}; border: 1px solid {COLORS['border']}; border-radius: 6px; padding: 4px; font-size: 12px; }} QMenu::item {{ padding: 8px 20px; border-radius: 4px; color: {COLORS['text_primary']}; }} QMenu::item:selected {{ background: {COLORS['accent_light']}; color: {COLORS['accent']}; }}")
        act_word = QAction("📝  Exportar como Word (.docx)", menu_exp)
        act_pdf  = QAction("📄  Exportar como PDF (.pdf)",  menu_exp)
        act_word.triggered.connect(self._exportar_word)
        act_pdf.triggered.connect(self._exportar_pdf)
        menu_exp.addAction(act_word); menu_exp.addAction(act_pdf)
        self.btn_exportar.setMenu(menu_exp)
        barra_topo_lay.addWidget(self.btn_exportar)

        lay_editor.addWidget(barra_topo)

        # Barra de formatação
        self.barra_fmt = BarraFormatacao(None)  # editor definido após criação
        lay_editor.addWidget(self.barra_fmt)

        # Info do documento
        self.lbl_info_doc = QLabel("")
        self.lbl_info_doc.setStyleSheet(f"color: {COLORS['text_muted']}; font-size: 10px; padding: 4px 16px;")
        lay_editor.addWidget(self.lbl_info_doc)

        # Editor de texto rico
        self.editor = QTextEdit()
        self.editor.setEnabled(False)
        self.editor.setPlaceholderText(
            "Selecione um documento na lista à esquerda para editar...\n\n"
            "Dica: use { } Variável para inserir campos automáticos como {{cliente}}, {{numero_processo}}, etc."
        )
        self.editor.setStyleSheet(f"""
            QTextEdit {{
                background: {COLORS['white']};
                border: none;
                padding: 24px 32px;
                font-family: 'Times New Roman', serif;
                font-size: 13px;
                color: {COLORS['text_primary']};
                line-height: 1.8;
            }}
        """)
        self.editor.textChanged.connect(self._marcar_modificado)
        lay_editor.addWidget(self.editor)

        # Conectar editor à barra de formatação
        self.barra_fmt.editor = self.editor

        raiz.addWidget(painel_editor)
        self.carregar_documentos()

    def carregar_documentos(self, filtro=""):
        cat_filtro = self.cb_filtro_cat.currentData() if hasattr(self, 'cb_filtro_cat') else ""
        self._dados = buscar_documentos(filtro)
        if cat_filtro:
            self._dados = [d for d in self._dados if d["categoria"] == cat_filtro]
        self.table.setRowCount(len(self._dados))
        for i, d in enumerate(self._dados):
            fav = "⭐ " if d.get("favorito") else ""
            self.table.setItem(i, 0, QTableWidgetItem(f"{fav}{d['titulo']}"))
            item_cat = QTableWidgetItem(d["categoria"][:8])
            item_cat.setForeground(QColor(COLORS["text_muted"]))
            self.table.setItem(i, 1, item_cat)

    def filtrar(self):
        self.carregar_documentos(self.busca.text())

    def _texto_com_marcacoes(self):
        """Extrai o texto do editor, reconstruindo {{IA:}} a partir do realce amarelo."""
        from PyQt6.QtGui import QTextCursor, QColor
        doc = self.editor.document()
        resultado = []
        bloco = doc.begin()
        while bloco.isValid():
            it = bloco.begin()
            while not it.atEnd():
                frag = it.fragment()
                if frag.isValid():
                    texto = frag.text()
                    fmt = frag.charFormat()
                    cor = fmt.background().color()
                    # Se o fundo é amarelo (realce IA), envolve em {{IA:}}
                    eh_ia = fmt.background().style() != 0 and cor.name().upper() in ("#FFF3A3",)
                    if eh_ia:
                        resultado.append("{{IA: " + texto + "}}")
                    else:
                        resultado.append(texto)
                it += 1
            resultado.append("\n")
            bloco = bloco.next()
        # Remove o último \n extra
        txt = "".join(resultado)
        if txt.endswith("\n"):
            txt = txt[:-1]
        return txt

    def _carregar_texto_editor(self, texto):
        """Carrega texto no editor. Se houver marcações {{IA:}}, aplica realce amarelo."""
        from core.ai.realce import tem_marcacoes, extrair_segmentos
        self.editor.blockSignals(True)
        if tem_marcacoes(texto):
            from PyQt6.QtGui import QTextCharFormat, QColor, QTextCursor
            self.editor.clear()
            cursor = self.editor.textCursor()
            fmt_normal = QTextCharFormat()
            fmt_ia = QTextCharFormat()
            fmt_ia.setBackground(QColor("#FFF3A3"))  # realce amarelo
            for trecho, eh_ia in extrair_segmentos(texto):
                cursor.insertText(trecho, fmt_ia if eh_ia else fmt_normal)
            self.editor.moveCursor(QTextCursor.MoveOperation.Start)
        else:
            self.editor.setPlainText(texto)
        self.editor.blockSignals(False)

    def abrir_documento(self):
        rows = self.table.selectionModel().selectedRows()
        if not rows: return
        if self._modificado:
            resp = QMessageBox.question(self, "Salvar?",
                f"O documento '{self._doc_aberto['titulo']}' foi modificado.\nSalvar antes de abrir outro?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel)
            if resp == QMessageBox.StandardButton.Yes: self.salvar_edicao()
            elif resp == QMessageBox.StandardButton.Cancel: return

        d = self._dados[rows[0].row()]
        self._doc_aberto = d
        self._modificado = False
        self._carregar_texto_editor(d.get("conteudo",""))
        self.editor.setEnabled(True)
        self.lbl_doc_aberto.setText(d["titulo"])
        data_br = "/".join(d["data_modificacao"].split("-")[::-1]) if "-" in d.get("data_modificacao","") else d.get("data_modificacao","")
        self.lbl_info_doc.setText(f"📁 {d['categoria']}  •  Modificado: {data_br}  •  Clique ⭐ para favoritar")
        self.btn_salvar.setEnabled(True)
        self.btn_exportar.setEnabled(True)
        self.btn_variaveis.setEnabled(True)
        self.btn_avaliar.setEnabled(True)
        self.lbl_status_salvo.setText("✔  Carregado")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['success']}; font-size: 11px; border: none;")

    def _marcar_modificado(self):
        if not self._doc_aberto: return
        self._modificado = True
        self.btn_salvar.setEnabled(True)
        self.lbl_status_salvo.setText("●  Não salvo")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['warning']}; font-size: 11px; font-weight: 600; border: none;")

    def _log_debug(self, msg):
        import os
        try:
            p = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "legis_debug.log")
            with open(p, "a", encoding="utf-8") as f:
                from datetime import datetime
                f.write(f"{datetime.now()} | {msg}\n")
                f.flush()
        except Exception:
            pass

    def salvar_edicao(self):
        if not self._doc_aberto: return
        self._log_debug("salvar_edicao: INICIO")
        try:
            self._log_debug("salvar_edicao: lendo editor")
            texto = self._texto_com_marcacoes()
            self._log_debug(f"salvar_edicao: texto lido ({len(texto)} chars)")
            dados = {"titulo": self._doc_aberto["titulo"],
                     "categoria": self._doc_aberto["categoria"],
                     "conteudo": texto}
            self._log_debug("salvar_edicao: chamando atualizar_documento")
            atualizar_documento(self._doc_aberto["id"], dados)
            self._log_debug("salvar_edicao: atualizar_documento OK")
        except Exception as e:
            self._log_debug(f"salvar_edicao: ERRO {e}")
            QMessageBox.critical(self, "Erro ao salvar", f"Não foi possível salvar:\n{e}")
            return
        self._doc_aberto["conteudo"] = dados["conteudo"]
        self._modificado = False
        self.btn_salvar.setEnabled(True)
        from datetime import datetime
        self.lbl_status_salvo.setText(f"✔  Salvo às {datetime.now().strftime('%H:%M')}")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['success']}; font-size: 11px; border: none;")
        self._log_debug("salvar_edicao: chamando carregar_documentos")
        self.carregar_documentos(self.busca.text())
        self._log_debug("salvar_edicao: FIM OK")

    def abrir_novo(self):
        dlg = NovoDocumentoDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            salvar_documento(dlg.obter_dados())
            self.carregar_documentos()
            self.table.selectRow(0)
            self.abrir_documento()

    def gerar_com_ia(self):
        from ui.gerar_peca_dialog import GerarPecaDialog
        dlg = GerarPecaDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            resultado = dlg.obter_resultado()
            if not resultado:
                return
            titulo, categoria, conteudo = resultado
            # Cria o documento e abre no editor
            salvar_documento({"titulo": titulo, "categoria": categoria, "conteudo": conteudo})
            self.carregar_documentos()
            self.table.selectRow(0)
            self.abrir_documento()
            QMessageBox.information(self, "Peça gerada",
                "A peça foi gerada e está aberta no editor.\n\n"
                "Revise o conteúdo, faça os ajustes necessários e use o botão "
                "Exportar para gerar o documento Word com formatação ABNT.")

    def avaliar_peca(self):
        if not self._doc_aberto:
            return
        from ui.avaliar_peca_dialog import AvaliarPecaDialog
        from core.database import (salvar_avaliacao, buscar_avaliacao_documento)
        existente = buscar_avaliacao_documento(self._doc_aberto["id"])
        dlg = AvaliarPecaDialog(self, titulo_peca=self._doc_aberto["titulo"],
                                avaliacao_existente=existente)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            res = dlg.obter_resultado()
            if not res:
                return
            conteudo = self.editor.toPlainText()
            salvar_avaliacao({
                "documento_id": self._doc_aberto["id"],
                "titulo": self._doc_aberto["titulo"],
                "tipo_peca": self._doc_aberto.get("categoria", ""),
                "area": self._doc_aberto.get("categoria", ""),
                "nota": res["nota"],
                "observacoes": res["observacoes"],
                "conteudo_resumo": conteudo[:1500],
            })
            QMessageBox.information(self, "Avaliação salva",
                f"Avaliação registrada ({res['nota']}★)!\n\n"
                "A IA vai considerar suas preferências nas próximas peças geradas.")

    def avaliar_peca(self):
        if not self._doc_aberto:
            return
        from ui.avaliar_peca_dialog import AvaliarPecaDialog
        from core.database import buscar_avaliacao_documento, salvar_avaliacao

        # Buscar avaliação existente
        aval = buscar_avaliacao_documento(self._doc_aberto["id"])
        nota_atual = aval.get("nota", 0) if aval else 0
        obs_atual = aval.get("observacoes", "") if aval else ""

        dlg = AvaliarPecaDialog(self, titulo_peca=self._doc_aberto["titulo"],
                                nota_atual=nota_atual, obs_atual=obs_atual)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            dados = dlg.obter_dados()
            # Extrair tipo/área do título (ex: "Petição Inicial - Direito Penal")
            titulo = self._doc_aberto["titulo"]
            tipo_peca = titulo.split(" - ")[0] if " - " in titulo else titulo
            area = titulo.split(" - ")[1] if " - " in titulo else ""
            salvar_avaliacao({
                "documento_id": self._doc_aberto["id"],
                "titulo": titulo,
                "categoria": self._doc_aberto["categoria"],
                "tipo_peca": tipo_peca,
                "area": area,
                "nota": dados["nota"],
                "observacoes": dados["observacoes"],
                "conteudo": self.editor.toPlainText(),
            })
            estrelas = "★" * dados["nota"] + "☆" * (5 - dados["nota"])
            QMessageBox.information(self, "Avaliação salva",
                f"Avaliação registrada: {estrelas}\n\n"
                "A IA usará isso para melhorar as próximas peças.")

    def importar_word(self):
        from ui.widgets import dialogo_abrir
        caminho = dialogo_abrir(self, "Importar documento Word", "Word (*.docx)")
        if not caminho: return

        self.lbl_status_salvo.setText("⏳  Importando...")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['blue']}; font-size: 11px; font-weight: 600; border: none;")

        if self._importador and self._importador.isRunning():
            self._importador.quit(); self._importador.wait()

        self._importador = ImportadorWord(caminho)
        self._importador.concluido.connect(self._importacao_concluida)
        self._importador.erro.connect(self._importacao_erro)
        self._importador.start()

    def _importacao_concluida(self, titulo, conteudo):
        # Pergunta categoria
        cats = CATEGORIAS
        cat, ok = QInputDialog.getItem(self, "Categoria do documento",
            "Selecione a categoria:", cats, 0, False)
        if not ok: cat = "Outro"
        salvar_documento({"titulo": titulo, "categoria": cat, "conteudo": conteudo})
        self.carregar_documentos()
        self.table.selectRow(0)
        self.abrir_documento()
        self.lbl_status_salvo.setText("✔  Documento importado com sucesso!")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['success']}; font-size: 11px; font-weight: 600; border: none;")

    def _importacao_erro(self, msg):
        self.lbl_status_salvo.setText("❌  Erro ao importar.")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['danger']}; font-size: 11px; font-weight: 600; border: none;")
        QMessageBox.critical(self, "Erro na importação",
            f"Não foi possível importar o arquivo.\n\nDetalhes:\n{msg[:300]}")

    def preencher_variaveis(self):
        if not self._doc_aberto: return
        dlg = PreencherVariaveisDialog(self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            numero = dlg.obter_numero()
            vars_ = buscar_variaveis_processo(numero)
            if not vars_:
                QMessageBox.warning(self, "Não encontrado",
                    "Processo não encontrado. Verifique o número informado."); return
            conteudo = self.editor.toPlainText()
            substituicoes = 0
            for var, valor in vars_.items():
                if var in conteudo:
                    conteudo = conteudo.replace(var, valor)
                    substituicoes += 1
            self.editor.blockSignals(True)
            self.editor.setPlainText(conteudo)
            self.editor.blockSignals(False)
            self._modificado = True
            self.btn_salvar.setEnabled(True)
            self.lbl_status_salvo.setText(f"⚡  {substituicoes} variável(is) preenchida(s) — salve para confirmar")
            self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['accent']}; font-size: 11px; font-weight: 600; border: none;")

    def excluir_selecionado(self):
        rows = self.table.selectionModel().selectedRows()
        if not rows:
            QMessageBox.information(self, "Selecione", "Selecione um documento."); return
        d = self._dados[rows[0].row()]
        resp = QMessageBox.question(self, "Confirmar",
            f"Excluir:\n{d['titulo']}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if resp == QMessageBox.StandardButton.Yes:
            excluir_documento(d["id"])
            self._doc_aberto = None; self._modificado = False
            self.editor.clear(); self.editor.setEnabled(False)
            self.lbl_doc_aberto.setText("Selecione ou crie um documento")
            self.lbl_info_doc.setText("")
            self.btn_salvar.setEnabled(False)
            self.btn_exportar.setEnabled(False)
            self.btn_variaveis.setEnabled(False)
            self.btn_avaliar.setEnabled(False)
            self.lbl_status_salvo.setText("")
            self.carregar_documentos()

    def _exportar_word(self):
        self.exportar("docx")

    def _exportar_pdf(self):
        self.exportar("pdf")

    def exportar(self, formato):
        # Proteção: verificar se o widget ainda é válido
        try:
            from PyQt6 import sip
            if sip.isdeleted(self):
                return
        except Exception:
            pass
        if not self._doc_aberto: return
        try:
            if self._modificado:
                resp = QMessageBox.question(self, "Salvar antes?",
                    "Há alterações não salvas. Salvar antes de exportar?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                if resp == QMessageBox.StandardButton.Yes: self.salvar_edicao()

            ext = "docx" if formato == "docx" else "pdf"
            nome = self._doc_aberto["titulo"].replace(" ", "_") + f".{ext}"

            # Pasta inicial segura: tenta Documentos, depois home (evita Desktop no OneDrive)
            pasta_inicial = os.path.expanduser("~")
            for cand in [os.path.join(pasta_inicial, "Documents"),
                         os.path.join(pasta_inicial, "Documentos")]:
                if os.path.isdir(cand):
                    pasta_inicial = cand
                    break

            # Usar diálogo NÃO-NATIVO do Qt — não depende do explorer do Windows,
            # que pode travar com pastas sincronizadas (OneDrive)
            dlg = QFileDialog(self, f"Exportar como {ext.upper()}")
            dlg.setOption(QFileDialog.Option.DontUseNativeDialog, True)
            dlg.setAcceptMode(QFileDialog.AcceptMode.AcceptSave)
            dlg.setDirectory(pasta_inicial)
            dlg.selectFile(nome)
            dlg.setNameFilter("Word (*.docx)" if formato == "docx" else "PDF (*.pdf)")
            if dlg.exec() != QFileDialog.DialogCode.Accepted:
                return
            selecionados = dlg.selectedFiles()
            if not selecionados:
                return
            caminho = selecionados[0]
            if not caminho.lower().endswith(f".{ext}"):
                caminho += f".{ext}"
        except Exception as e:
            QMessageBox.critical(self, "Erro", f"Não foi possível abrir o diálogo de exportação:\n{e}")
            return

        self.lbl_status_salvo.setText("⏳  Exportando...")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['blue']}; font-size: 11px; font-weight: 600; border: none;")

        if self._exportador and self._exportador.isRunning():
            return  # já exportando

        self._exportador = ExportadorThread(formato, caminho,
            self._doc_aberto["titulo"], self._doc_aberto["categoria"],
            self.editor.toPlainText())
        self._exportador.concluido.connect(self._exportacao_concluida)
        self._exportador.erro.connect(self._exportacao_erro)
        self._exportador.start()

    def _exportacao_concluida(self, caminho):
        self.lbl_status_salvo.setText("✔  Exportado!")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['success']}; font-size: 11px; font-weight: 600; border: none;")
        resp = QMessageBox.question(self, "Exportado",
            f"Arquivo salvo em:\n{caminho}\n\nAbrir agora?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if resp == QMessageBox.StandardButton.Yes:
            try:
                if os.name == "nt":
                    # os.startfile pode causar crash nativo em pastas sincronizadas;
                    # usar subprocess com 'start' é mais seguro
                    subprocess.Popen(["cmd", "/c", "start", "", caminho], shell=False,
                                     creationflags=0x08000000)
                else:
                    subprocess.Popen(["xdg-open", caminho])
            except Exception:
                QMessageBox.information(self, "Abrir manualmente",
                    f"O arquivo foi salvo com sucesso.\n\nAbra manualmente em:\n{caminho}")

    def _exportacao_erro(self, msg):
        self.lbl_status_salvo.setText("❌  Erro ao exportar.")
        self.lbl_status_salvo.setStyleSheet(f"color: {COLORS['danger']}; font-size: 11px; font-weight: 600; border: none;")
        QMessageBox.critical(self, "Erro", f"Não foi possível exportar:\n{msg[:300]}")
