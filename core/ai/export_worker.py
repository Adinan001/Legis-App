# core/ai/export_worker.py
# Worker de exportação de documentos (docx/pdf) em processo separado,
# para isolar python-docx/reportlab do Qt e evitar crashes nativos.
import sys, os, json
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))


def exportar_docx(caminho, titulo, categoria, conteudo):
    import re
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(3); section.bottom_margin = Cm(2)
        section.left_margin = Cm(3); section.right_margin = Cm(2)

    PADRAO_IA = re.compile(r"\{\{IA:\s*(.*?)\}\}", re.DOTALL)

    def add_runs(par, texto, realce_ia=False):
        """Adiciona runs tratando **negrito**, com opção de realce amarelo."""
        partes = re.split(r"(\*\*.*?\*\*)", texto)
        for parte in partes:
            if not parte:
                continue
            if parte.startswith("**") and parte.endswith("**") and len(parte) > 4:
                run = par.add_run(parte[2:-2]); run.font.bold = True
            else:
                run = par.add_run(parte)
            run.font.size = Pt(12); run.font.name = "Times New Roman"
            if realce_ia:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    doc.add_paragraph()
    _lixo = re.compile(r"^\(?\s*(linha em branco|em branco|centralizado|em negrito|espa[çc]o)\s*\)?\s*$", re.IGNORECASE)
    for linha in conteudo.split("\n"):
        linha = re.sub(r"^#{1,6}\s*", "", linha)
        if re.fullmatch(r"[-*_=]{3,}", (linha.strip() or "x")):
            continue
        if _lixo.match(linha.strip()):
            continue
        par = doc.add_paragraph()
        # Para detectar título, usar a linha sem marcações IA
        linha_sem_ia = PADRAO_IA.sub(r"\1", linha).strip()
        eh_titulo = linha_sem_ia.startswith("**") and linha_sem_ia.endswith("**") and linha_sem_ia.count("**") == 2
        if eh_titulo:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if any(
                k in linha_sem_ia.upper() for k in ["EXCELENT", "EGRÉGIO", "COLENDO", "SENHOR"]) else WD_ALIGN_PARAGRAPH.LEFT
        else:
            par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not eh_titulo and linha_sem_ia:
            par.paragraph_format.first_line_indent = Cm(2)

        # Separar segmentos IA (realce) e normais
        pos = 0
        for m in PADRAO_IA.finditer(linha):
            if m.start() > pos:
                add_runs(par, linha[pos:m.start()], realce_ia=False)
            add_runs(par, m.group(1), realce_ia=True)
            pos = m.end()
        if pos < len(linha):
            add_runs(par, linha[pos:], realce_ia=False)

    doc.save(caminho)


def exportar_pdf(caminho, titulo, categoria, conteudo):
    import re
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER

    doc = SimpleDocTemplate(caminho, pagesize=A4,
        topMargin=3*cm, bottomMargin=2*cm, leftMargin=3*cm, rightMargin=2*cm)
    estilos = getSampleStyleSheet()
    normal = ParagraphStyle("just", parent=estilos["Normal"],
        fontName="Times-Roman", fontSize=12, leading=18, alignment=TA_JUSTIFY,
        firstLineIndent=2*cm)
    centro = ParagraphStyle("centro", parent=estilos["Normal"],
        fontName="Times-Bold", fontSize=12, leading=18, alignment=TA_CENTER)

    elementos = [Spacer(1, 12)]
    _lixo = re.compile(r"^\(?\s*(linha em branco|em branco|centralizado|em negrito|espa[çc]o)\s*\)?\s*$", re.IGNORECASE)
    for linha in conteudo.split("\n"):
        linha = re.sub(r"^#{1,6}\s*", "", linha)
        if re.fullmatch(r"[-*_=]{3,}", (linha.strip() or "x")):
            continue
        if _lixo.match(linha.strip()):
            continue
        if not linha.strip():
            elementos.append(Spacer(1, 8)); continue
        txt = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", linha)
        linha_strip = linha.strip()
        eh_titulo = linha_strip.startswith("**") and linha_strip.endswith("**")
        try:
            elementos.append(Paragraph(txt, centro if eh_titulo else normal))
        except Exception:
            elementos.append(Paragraph(linha.replace("&","&amp;").replace("<","&lt;"), normal))
    doc.build(elementos)


def main():
    try:
        entrada = json.loads(sys.stdin.read())
    except Exception as e:
        print(json.dumps({"ok": False, "erro": f"Entrada inválida: {e}"})); return

    formato = entrada.get("formato", "docx")
    caminho = entrada.get("caminho", "")
    titulo = entrada.get("titulo", "")
    categoria = entrada.get("categoria", "")
    conteudo = entrada.get("conteudo", "")

    try:
        import tempfile, shutil
        ext = ".docx" if formato == "docx" else ".pdf"
        fd, tmp = tempfile.mkstemp(suffix=ext); os.close(fd)
        if formato == "docx":
            exportar_docx(tmp, titulo, categoria, conteudo)
        else:
            exportar_pdf(tmp, titulo, categoria, conteudo)
        os.makedirs(os.path.dirname(caminho), exist_ok=True)
        shutil.move(tmp, caminho)
        print(json.dumps({"ok": True, "caminho": caminho}))
    except Exception as e:
        import traceback
        print(json.dumps({"ok": False, "erro": str(e), "trace": traceback.format_exc()[:400]}))


if __name__ == "__main__":
    main()
